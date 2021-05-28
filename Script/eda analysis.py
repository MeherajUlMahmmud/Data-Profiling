#%%
import numpy as np

import pandas as pd

import matplotlib.pyplot as plt
from matplotlib.patches import Rectangle

import seaborn as sns



#%%
name = "abalone.csv"
try:
    df = pd.read_csv(name)
except:
    print("Error: Data file not found!")

# print(name[:-4])
# df.columns = ['crim', 'zn', 'indus', 'chas', 'nox', 'rm', 'age', 'dis', 'rad', 'tax', 'ptratio', 'b', 'lstat', 'medv']
# print(df.head())
rows = len(df.axes[0])
cols = len(df.axes[1])
print(rows, cols)

#%%
raw_num_df = df.describe()
print(raw_num_df)

#%%
no_of_rows = len(df.columns)

data_qlt_df = pd.DataFrame(index=np.arange(0, no_of_rows),
                            columns=('column_name', 'col_data_type','non_null_values',
                                     'unique_values_count', 'column_dtype')
                          )

for i, cols in enumerate(df.columns):
    col_unique_count = df[cols].nunique()
    
    data_qlt_df.loc[i] = [cols,
                          df[cols].dtype,
                          df[cols].count(),
                          col_unique_count,
                          cols + '~'+ str(df[cols].dtype)
                          ]

#%%
raw_num_df = df.describe().T.round(2)


#%%
data_qlt_df = pd.merge(data_qlt_df, raw_num_df, how='left', left_on='column_name', right_index=True)


#%%
data_qlt_df['%_of_non_nulls'] = (data_qlt_df['non_null_values'] / df.shape[0]) * 100

data_qlt_df['null_values'] = df.shape[0] - data_qlt_df['non_null_values']

data_qlt_df['%_of_nulls'] = 100 - data_qlt_df['%_of_non_nulls']

data_qlt_df["dtype_count"] = data_qlt_df.groupby('col_data_type')["col_data_type"].transform('count')

data_qlt_df["count"] = data_qlt_df['null_values'] + data_qlt_df['non_null_values']


#%%
data_qlt_df = data_qlt_df[
                            ['column_name', 'col_data_type', 'non_null_values', '%_of_non_nulls',
                             'null_values', '%_of_nulls', 'unique_values_count',
                             'count', 'mean', 'std', 'min', '25%',
                             '50%', '75%', 'max']
                         ]
    
    
# data_qlt_df = data_qlt_df[['column_name', 'col_data_type', 'non_null_values', '%_of_non_nulls',
#                              'null_values', '%_of_nulls', 'unique_values_count', 'count', 'mean',
#                              'min', '25%', '50%', '75%', 'max']]


#%%
def missing_data_count(data):
    total = data.isnull().sum().sort_values(ascending=False)
    percent = (data.isnull().sum()/data.isnull().count()).sort_values(ascending=False)
    missing_data = pd.concat([total, percent], axis=1, keys=['Total', 'Percent'])
    return missing_data

print(missing_data_count(df))

#%%
from pandas.api.types import is_numeric_dtype
from pandas.api.types import is_string_dtype

num_cols = [cols for cols in df.columns if is_numeric_dtype(df[cols]) and len(df[cols].dropna()) > 0]
cat_cols = [cols for cols in df.columns if is_string_dtype(df[cols]) and len(df[cols].dropna()) > 0]
print(num_cols)
print(cat_cols)


#%%

for col_name in num_cols:
    
    no_null_col = df[col_name].dropna()
    
    
    # Plot the graphs
    fig3 = plt.figure(figsize=(20, 15))
    fig3.suptitle("Profile of " + col_name, fontsize=25)
    plt.subplots_adjust(wspace=0.4, hspace=0.35)



    ax1 = fig3.add_subplot(2, 2, 1)
    ax1.set_title("Box Plot", fontsize=20)
    plt.setp(ax1.get_xticklabels(), fontsize=10)
    plt.setp(ax1.get_yticklabels(), fontsize=15)
    ax1.boxplot(no_null_col)



    ax1 = fig3.add_subplot(2, 2, 2)
    ax1.set_title("Data Distribution", fontsize=20)
    plt.setp(ax1.get_xticklabels(), fontsize=10)
    plt.setp(ax1.get_yticklabels(), fontsize=15)
    ax1.hist(no_null_col)



    ax1 = fig3.add_subplot(2, 2, 3)
    ax1.set_title("Box plot without outliers", fontsize=20)
    plt.setp(ax1.get_xticklabels(), fontsize=10)
    plt.setp(ax1.get_yticklabels(), fontsize=15)
    ax1.boxplot(no_null_col, showfliers=False)




    fig3.suptitle("Profile of column  " + col_name, fontsize=25)  #Title for the whole figure
    fig_name = 'fig_' + col_name
    fig3.savefig(fig_name, dpi=50)
    plt.close('all')
    
    print(col_name, "    completed")
    

#%%
# for cat in cat_cols:
#     print(cat, df[cat].unique(), df[cat].nunique())
#     print()
    

#%%

for col_name in cat_cols:
    print(col_name, " completed")
    
    no_null_col = df[col_name].dropna()

    values_freq_threshold = 25
    col_unique_count = df[col_name].nunique()
    
    col_unique_vals = df[col_name].value_counts(normalize=True, sort=True)

    
    fig4 = plt.figure(figsize=(20,7))
    fig4.suptitle("Profile of column  " + col_name, fontsize=25)
    # plt.subplots_adjust(wspace=0.4, hspace=0.35, bottom=0.35)

    ax1 = fig4.add_subplot(1,1,1)
    plt.setp(ax1.get_xticklabels(), fontsize=12)
    plt.setp(ax1.get_yticklabels(), fontsize=12)
    
    col_unique_vals.head(values_freq_threshold).sort_values(ascending=False).plot.bar()
    plt.xticks(rotation=0)
    for p in ax1.patches:
        ax1.annotate(str(round(p.get_height(),2)), (p.get_x() * 1.005, p.get_height() * 1.005), fontsize=15)
    
    fig4.suptitle("Profile of column  " + col_name)
    fig_name = 'fig_' + col_name
    fig4.savefig(fig_name, dpi= 50)

    plt.close('all')


#%%
fig, ax = plt.subplots(figsize=(15, 10))
plt.subplots_adjust(bottom=0.35)
plt.autoscale()

corr_data = df.corr()

sns.heatmap(corr_data, annot=True)

fig_name = 'fig_cor_plot.png'
fig.savefig(fig_name,  dpi=70)
plt.close('all')


#%%
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


#%%
document = Document()
doc_style = 'Table Grid'
document.add_heading('Exploratory Data Analysis', 0)


#%%
document.add_page_break()
document.add_heading('Columns Data Profile Summary', 0)



#%%

p = document.add_paragraph(' ')

# Heading 1
document.add_heading('Dataset shape', level=1)

table = document.add_table(rows=2, cols=2, style = doc_style)

# Header row
cell = table.cell(0, 0)
cell.text = 'No.of rows'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = True

cell = table.cell(0, 1)
cell.text = 'No.of columns'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = True

# Values
cell = table.cell(1, 0)
cell.text = F'{df.shape[0] :,}'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False

cell = table.cell(1, 1)
cell.text = F'{df.shape[1] :,}'
cell_font = cell.paragraphs[0].runs[0].font
cell_font.size = Pt(11)
cell_font.bold = False



#%%

p = document.add_paragraph(' ')

# Heading 1
document.add_heading('Dataframe columns summary', level=1)

# Rehsape the column data type dataframe into form that can be printed in MS Word
data = round(data_qlt_df[['column_name','col_data_type', 'non_null_values', 'null_values', 'count']], 2)

# add a table to the end and create a reference variable
# extra row is so we can add the header row
table = document.add_table(data.shape[0]+1, data.shape[1], style=doc_style)

# add the header rows.
for j in range(data.shape[1]):

    #header row first two columns
    if j <= 1:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
    else:
        cell = table.cell(0, j)
        cell.text = F'{data.columns[j]}'
        cell_font = cell.paragraphs[0].runs[0].font
        cell_font.size = Pt(11)
        cell_font.bold = True
        cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT
        
    
# add the rest of the data frame
for i in range(data.shape[0]):
    for j in range(data.shape[1]):
        if j <= 1:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j]}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False            
        else:
            cell = table.cell(i+1, j)
            cell.text = F'{data.values[i,j] :,}'
            cell_font = cell.paragraphs[0].runs[0].font
            cell_font.size = Pt(11)
            cell_font.bold = False
            cell.paragraphs[0].alignment= WD_ALIGN_PARAGRAPH.RIGHT


#%%
document.add_page_break()
document.add_heading('Data correlation plot', 0)

p = document.add_paragraph('')

document.add_picture('fig_cor_plot.png', height=Inches(6), width=Inches(7))

#%%
document.add_page_break()
document.add_heading('Column Details', 0)


#%%

for i in range(data_qlt_df.shape[0]):
    document.add_page_break()
    
    table = document.add_table(rows=4, cols=3, style=doc_style )

    """
    ADD VALUES TO TABLE
    """


    # Cell 0, 0: Column name
    cell = table.cell(0, 0)
    cell.text = data_qlt_df["column_name"][i]
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(15)
    cell_font.bold = True

    # Cell 0, 1: Column data type
    cell = table.cell(0, 1)
    cell.text = 'Data Type : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(str(data_qlt_df["col_data_type"][i]))
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(12)
    cell_font2.bold = False

    # Cell 0, 2: Count of toal values in the column
    cell = table.cell(0, 2)
    cell.text = 'Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["count"][i] :,.0f}')
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 1, 0: Count of unique values in the column
    cell = table.cell(1, 0)
    cell.text = 'Unique Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    unique_per = (data_qlt_df["unique_values_count"][i] / data_qlt_df["count"][i]) * 100
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["unique_values_count"][i] :,.0f}' + "   " + F'({unique_per :,.2f}%)' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 1, 1: Count of non-null values in the column
    cell = table.cell(1, 1)
    cell.text = 'Non-Null Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["non_null_values"][i] :,.0f}' + "   " + F' ({data_qlt_df["%_of_non_nulls"][i]  :,.2f}%)' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False       

    # Cell 1, 2: Count of null values in the column
    cell = table.cell(1, 2)
    cell.text = 'Null Values Count : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["null_values"][i]  :,.0f}' + "   " + F' ({data_qlt_df["%_of_nulls"][i]  :,.2f}%)' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 2, 0: Min of values in the column
    cell = table.cell(2, 0)
    cell.text = 'Min : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["min"][i]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 2, 1: Mean of values in the column
    cell = table.cell(2, 1)
    cell.text = 'Mean :  \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["mean"][i] :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 2,2: Max of values in the column
    cell = table.cell(2, 2)
    cell.text = 'Max : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["max"][i]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 3, 0: 25th Percentile of values in the column
    cell = table.cell(3, 0)
    cell.text = '25th Percentile : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["25%"][i]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 3, 1: 50th Percentile of values in the column
    cell = table.cell(3, 1)
    cell.text = '50th Percentile : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["50%"][i]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False

    # Cell 3, 2: 75th Percentile of values in the column
    cell = table.cell(3, 2)
    cell.text = '75th Percentile : \n'
    cell_font = cell.paragraphs[0].runs[0].font
    cell_font.size = Pt(11)
    cell_font.bold = True
    p = cell.paragraphs[0].add_run(F'{data_qlt_df["75%"][i]  :,.2f}' )
    cell_font2 = cell.paragraphs[0].runs[1].font
    cell_font2.size = Pt(11)
    cell_font2.bold = False


    p = document.add_paragraph(' ')
    p = document.add_paragraph(' ')

    fig_name = 'fig_' + data_qlt_df['column_name'][i] + '.png'
    document.add_picture(fig_name, height=Inches(3.5), width=Inches(6))


#%%
document.save(name[:-4] + '.docx')



